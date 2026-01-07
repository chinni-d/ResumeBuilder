from flask import make_response
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from io import BytesIO

def get_template_colors(template_id):
    """Get template-specific colors and styling"""
    templates = {
        1: {'primary': '#333333', 'secondary': '#666666', 'name': 'Classic'},
        2: {'primary': '#3498db', 'secondary': '#2980b9', 'name': 'Modern Blue'},
        3: {'primary': '#27ae60', 'secondary': '#2ecc71', 'name': 'Creative Green'},
        4: {'primary': '#8e44ad', 'secondary': '#9b59b6', 'name': 'Executive Purple'},
        5: {'primary': '#e67e22', 'secondary': '#f39c12', 'name': 'Tech Orange'},
        6: {'primary': '#7f8c8d', 'secondary': '#95a5a6', 'name': 'Minimal Gray'},
        7: {'primary': '#e74c3c', 'secondary': '#c0392b', 'name': 'Bold Red'},
        8: {'primary': '#17a2b8', 'secondary': '#138496', 'name': 'Elegant Teal'}
    }
    return templates.get(template_id, templates[1])

def export_to_docx(resume):
    """Export resume to DOCX format with template styling"""
    try:
        doc = Document()
        
        # Get template colors
        template = get_template_colors(resume.template_id)
        primary_rgb = tuple(int(template['primary'][i:i+2], 16) for i in (1, 3, 5))
        secondary_rgb = tuple(int(template['secondary'][i:i+2], 16) for i in (1, 3, 5))
        
        # Add content to document
        if resume.content:
            content = resume.content
            
            # Header with name and styling based on template
            if content.full_name:
                if resume.template_id in [2, 3, 5, 7]:  # Templates with colored headers
                    # Create a header paragraph with background color
                    header = doc.add_heading(content.full_name, 0)
                    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Set font color to white for colored backgrounds
                    for run in header.runs:
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.size = Inches(0.3)
                    
                    # Add background color to paragraph (this is complex in python-docx)
                    p = header._element
                    pPr = p.get_or_add_pPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:fill'), template['primary'][1:])  # Remove # from color
                    pPr.append(shd)
                else:
                    header = doc.add_heading(content.full_name, 0)
                    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in header.runs:
                        run.font.color.rgb = RGBColor(*primary_rgb)
            
            # Contact information
            contact_info = []
            if content.email: contact_info.append(f"Email: {content.email}")
            if content.phone: contact_info.append(f"Phone: {content.phone}")
            if content.address: contact_info.append(f"Address: {content.address}")
            if content.linkedin: contact_info.append(f"LinkedIn: {content.linkedin}")
            if content.github: contact_info.append(f"GitHub: {content.github}")
            if content.website: contact_info.append(f"Website: {content.website}")
            
            if contact_info:
                contact_para = doc.add_paragraph(" | ".join(contact_info))
                contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in contact_para.runs:
                    run.font.color.rgb = RGBColor(*secondary_rgb)
                    run.font.size = Inches(0.12)
                doc.add_paragraph()  # Empty line
            
            # Professional Summary
            if content.summary:
                summary_heading = doc.add_heading('Professional Summary', level=1)
                for run in summary_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                doc.add_paragraph(content.summary)
                doc.add_paragraph()
            
            # Skills with ratings
            skills = content.get_skills()
            if skills:
                skills_heading = doc.add_heading('Skills', level=1)
                for run in skills_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                
                # Create skills table for better formatting
                skills_table = doc.add_table(rows=0, cols=2)
                skills_table.style = 'Light Grid Accent 1'
                
                for skill in skills:
                    row = skills_table.add_row()
                    row.cells[0].text = skill.get('name', '')
                    stars = '★' * skill.get('level', 3) + '☆' * (5 - skill.get('level', 3))
                    row.cells[1].text = stars
                    
                    # Color the skill name
                    for paragraph in row.cells[0].paragraphs:
                        for run in paragraph.runs:
                            run.font.color.rgb = RGBColor(*secondary_rgb)
                            run.font.bold = True
                
                doc.add_paragraph()
            
            # Experience
            experience = content.get_experience()
            if experience:
                exp_heading = doc.add_heading('Professional Experience', level=1)
                for run in exp_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                
                for exp in experience:
                    if exp.get('title'):
                        exp_title = exp['title']
                        if exp.get('company'):
                            exp_title += f" - {exp['company']}"
                        if exp.get('period'):
                            exp_title += f" ({exp['period']})"
                        
                        exp_para = doc.add_heading(exp_title, level=2)
                        for run in exp_para.runs:
                            run.font.color.rgb = RGBColor(*secondary_rgb)
                        
                        if exp.get('description'):
                            doc.add_paragraph(exp['description'])
                doc.add_paragraph()
            
            # Education
            education = content.get_education()
            if education:
                edu_heading = doc.add_heading('Education', level=1)
                for run in edu_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                
                for edu in education:
                    if edu.get('degree'):
                        edu_text = edu['degree']
                        if edu.get('school'):
                            edu_text += f" - {edu['school']}"
                        if edu.get('year'):
                            edu_text += f" ({edu['year']})"
                        if edu.get('gpa'):
                            edu_text += f" - GPA: {edu['gpa']}"
                        
                        edu_para = doc.add_paragraph(edu_text)
                        for run in edu_para.runs:
                            run.font.bold = True
                doc.add_paragraph()
            
            # Projects
            projects = content.get_projects()
            if projects:
                proj_heading = doc.add_heading('Projects', level=1)
                for run in proj_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                
                for proj in projects:
                    if proj.get('name'):
                        proj_name = doc.add_heading(proj['name'], level=2)
                        for run in proj_name.runs:
                            run.font.color.rgb = RGBColor(*secondary_rgb)
                        
                        if proj.get('description'):
                            doc.add_paragraph(proj['description'])
                        if proj.get('technologies'):
                            tech_para = doc.add_paragraph(f"Technologies: {proj['technologies']}")
                            for run in tech_para.runs:
                                run.font.italic = True
                        if proj.get('link'):
                            link_para = doc.add_paragraph(f"Link: {proj['link']}")
                            for run in link_para.runs:
                                run.font.color.rgb = RGBColor(*primary_rgb)
                doc.add_paragraph()
            
            # Certifications
            certifications = content.get_certifications()
            if certifications:
                cert_heading = doc.add_heading('Certifications', level=1)
                for run in cert_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                
                for cert in certifications:
                    if cert.get('name'):
                        cert_text = cert['name']
                        if cert.get('issuer'):
                            cert_text += f" - {cert['issuer']}"
                        if cert.get('date'):
                            cert_text += f" ({cert['date']})"
                        
                        cert_para = doc.add_paragraph(cert_text)
                        for run in cert_para.runs:
                            run.font.bold = True
                doc.add_paragraph()
            
            # Awards
            awards = content.get_awards()
            if awards:
                award_heading = doc.add_heading('Awards & Achievements', level=1)
                for run in award_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                
                for award in awards:
                    if award.get('name'):
                        award_text = award['name']
                        if award.get('date'):
                            award_text += f" ({award['date']})"
                        
                        award_name = doc.add_heading(award_text, level=2)
                        for run in award_name.runs:
                            run.font.color.rgb = RGBColor(*secondary_rgb)
                        
                        if award.get('description'):
                            doc.add_paragraph(award['description'])
                doc.add_paragraph()
            
            # Languages
            languages = content.get_languages()
            if languages:
                lang_heading = doc.add_heading('Languages', level=1)
                for run in lang_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                
                lang_items = [f"{lang.get('name', '')} ({lang.get('level', 'Intermediate')})" for lang in languages]
                doc.add_paragraph(", ".join(lang_items))
                doc.add_paragraph()
            
            # Hobbies
            hobbies = content.get_hobbies()
            if hobbies:
                hobby_heading = doc.add_heading('Hobbies & Interests', level=1)
                for run in hobby_heading.runs:
                    run.font.color.rgb = RGBColor(*primary_rgb)
                doc.add_paragraph(", ".join(hobbies))
        
        # Save to BytesIO
        docx_file = BytesIO()
        doc.save(docx_file)
        docx_file.seek(0)
        
        response = make_response(docx_file.read())
        response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        response.headers['Content-Disposition'] = f'attachment; filename="{resume.title}.docx"'
        
        return response
    except Exception as e:
        print(f"DOCX export error: {e}")
        return make_response("Error generating DOCX", 500)